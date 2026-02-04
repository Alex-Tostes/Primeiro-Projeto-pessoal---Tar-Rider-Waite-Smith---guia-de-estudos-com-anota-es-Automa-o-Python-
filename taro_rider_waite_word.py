from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# Caminho da pasta com as imagens renomeadas
pasta = r"C:\Users\SYS\Documents\Cards-png\Cards-png"

# Dados das cartas: [nome, palavras-chave, descrição detalhada]
cartas_info = [
    ["The Fool", "Inocência, espontaneidade, novos começos / aventura, liberdade, confiança", 
     "Representa novas jornadas, coragem e espírito livre, pronto para explorar o desconhecido."],
    ["The Magician", "Habilidade, concentração, poder, ação / criatividade, destreza, manifestação", 
     "Simboliza poder de transformação, foco e habilidade para criar a realidade desejada."],
    ["High Priestess", "Intuição, mistério, sabedoria interior / segredo, espiritualidade, insight", 
     "Convida à introspecção, à confiança na intuição e à sabedoria oculta."],
    ["The Empress", "Fertilidade, abundância, cuidado, criatividade / nutrição, crescimento, prosperidade", 
     "Carta da abundância, criação e conexão com a natureza e as emoções."],
    ["The Emperor", "Autoridade, estabilidade, liderança, proteção / estrutura, disciplina, segurança", 
     "Representa ordem, liderança e poder para estruturar a vida e alcançar objetivos."],
    ["The Hierophant", "Tradição, espiritualidade, ensino, conformidade / orientação, sabedoria, ritual", 
     "Simboliza ensino, tradição e aprendizado dentro de um sistema de valores estabelecido."],
    ["The Lovers", "Amor, união, parcerias, escolhas / harmonia, compromisso, conexão", 
     "Fala de relacionamentos profundos, escolhas significativas e alinhamento de valores."],
    ["The Chariot", "Vitória, determinação, controle, disciplina / força de vontade, conquista, foco", 
     "Indica superação de obstáculos por meio de foco, disciplina e ação determinada."],
    ["Strength", "Coragem, força interior, paciência, compaixão / autocontrole, resiliência, coragem", 
     "Força interior, domínio sobre impulsos e coragem para enfrentar desafios com compaixão."],
    ["The Hermit", "Reflexão, introspecção, busca interior / sabedoria, paciência, orientação", 
     "Momento de reflexão profunda, busca de conhecimento interno e iluminação pessoal."],
    ["Wheel of Fortune", "Destino, ciclos, mudança, oportunidade / sorte, evolução, transformação", 
     "Ciclos da vida, mudanças inevitáveis e oportunidades que surgem com a rotação do destino."],
    ["Justice", "Justiça, verdade, equilíbrio, responsabilidade / ética, decisão, imparcialidade", 
     "Decisões importantes, consequências das ações e equilíbrio entre mente e coração."],
    ["The Hanged Man", "Sacrifício, espera, perspectiva, renúncia / paciência, entrega, visão", 
     "Necessidade de olhar a situação por outro ângulo e aceitar pausas temporárias."],
    ["Death", "Fim de ciclo, transformação, renascimento, libertação / mudança profunda, transição", 
     "Simboliza término e renascimento, transformação e deixar o antigo para abrir espaço ao novo."],
    ["Temperance", "Equilíbrio, moderação, propósito, paciência / harmonia, mistura, adaptação", 
     "Equilíbrio emocional e espiritual, moderação e integração de opostos."],
    ["The Devil", "Tentação, vícios, materialismo, limitação / medo, apego, obsessão", 
     "Aponta para correntes internas ou externas que limitam o crescimento e a liberdade."],
    ["The Tower", "Destruição, despertar, revelação, choque / ruptura, mudança repentina", 
     "Mudanças inesperadas e profundas que derrubam estruturas antigas, abrindo caminho para renovação."],
    ["The Star", "Esperança, inspiração, serenidade, orientação / cura, clareza, luz interior", 
     "Renovação, esperança e confiança no futuro, inspiração e cura interior."],
    ["The Moon", "Ilusão, medo, intuição, ciclos / mistério, engano, subconsciente", 
     "Exploração do inconsciente, ilusões e necessidade de confiar na intuição."],
    ["The Sun", "Sucesso, alegria, vitalidade, realização / clareza, energia positiva", 
     "Felicidade, sucesso e clareza, iluminando o caminho com vitalidade e alegria."],
    ["Judgement", "Renascimento, julgamento, chamado, decisão / avaliação, despertar, transformação", 
     "Convite à avaliação de escolhas passadas e despertar para uma nova fase."],
    ["The World", "Conclusão, realização, viagem, plenitude / integração, sucesso, celebração", 
     "Culminação de um ciclo, realização plena e integração de experiências."],
]

# Função para gerar o documento
def gerar_word():
    arquivos = sorted(os.listdir(pasta))
    doc = Document()

    carta_index = 0  # Para percorrer apenas cartas válidas

    for arquivo in arquivos:
        img_path = os.path.join(pasta, arquivo)
        if "_" not in arquivo or carta_index >= len(cartas_info):
            continue  # ignora arquivos inválidos

        if carta_index != 0:
            doc.add_page_break()

        # Adiciona imagem maior
        try:
            doc.add_picture(img_path, width=Inches(5))  # aumenta imagem
        except Exception as e:
            print(f"Erro ao adicionar imagem {arquivo}: {e}")
            continue

        nome, palavras, descricao = cartas_info[carta_index]

        # Nome da carta
        p_nome = doc.add_paragraph()
        p_nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_nome = p_nome.add_run(nome)
        run_nome.bold = True
        run_nome.font.size = Pt(18)

        # Palavras-chave
        p_kw = doc.add_paragraph()
        p_kw.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_kw = p_kw.add_run(palavras)
        run_kw.font.size = Pt(12)
        run_kw.italic = True

        # Descrição detalhada
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run_desc = p_desc.add_run(descricao)
        run_desc.font.size = Pt(12)

        # Espaço maior para anotações
        doc.add_paragraph("\n" * 12 + "Minhas Notas: ____________________________________________\n" * 8)

        carta_index += 1

    # Salva documento
    doc.save(r"C:\Users\SYS\Documents\Taro_Rider_Waite_Completo.docx")
    print("Word completo gerado com sucesso em C:\\Users\\SYS\\Documents\\Taro_Rider_Waite_Completo.docx")

# Executa
gerar_word()
